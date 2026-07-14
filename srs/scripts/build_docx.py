#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the GeoE3 SRS docx by writing directly into the Kartoza template.

No pandoc: the Kartoza_Document_Template.docx is opened with python-docx, its
demo pages are harvested for exemplar formatting (default table, requirements
keyword colours, bullet/numbered list numbering, figure captions) and then
removed, and the SRS markdown is rendered straight into the document using
the template's own styles. Finally the cover artwork is swapped and the
header/metadata placeholders are filled at zip level.

Usage: build_docx.py <template.docx> <srs.md> <cover.png> <out.docx>
"""

import copy
import re
import struct
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor
from PIL import Image

TEMPLATE, SRS_MD, COVER, OUT = (Path(a) for a in sys.argv[1:5])
BASE = SRS_MD.parent

CHARCOAL = RGBColor(0x38, 0x39, 0x39)
GREY = RGBColor(0x67, 0x68, 0x69)
AMBER_TEXT = RGBColor(0xB2, 0x6B, 0x00)  # readable amber for (SHOULD)

MAX_W_IN, MAX_H_IN = 6.25, 7.6

REPLACEMENTS = {
    "{Project name}": "GeoE3 — Geospatial Enabling Environments for Employment",
    "Kartoza (Pty) Ltd &amp; Kartoza Lda": "The World Bank · developed by Kartoza",
    "{Type}": "Software Requirements Specification (SRS)",
    "{0.1}": "1.0.0",
    "{Name}": "Tim Sutton",
    "{Draft / In review / Final}": "Draft",
    "{Month Year}": "July 2026",
    "{DOCUMENT TITLE}": "GEOE3",
    "{SUBTITLE}": "SRS draft v1.0.0",
}

doc = Document(str(TEMPLATE))
body = doc.element.body


# --------------------------------------------------------------------------
# 1. Harvest exemplars from the demo pages
# --------------------------------------------------------------------------
def paragraph_text(p):
    return "".join(node.text or "" for node in p.iter(qn("w:t")))


demo_start = None
bullet_numpr = None
ordered_numpr = None
caption_ppr = None
caption_rpr = None
must_rpr = None
should_rpr = None
may_rpr = None
exemplar_tbl = None

paragraphs = body.findall(qn("w:p"))
for p in paragraphs:
    text = paragraph_text(p)
    if demo_start is None and "How to use this template" in text:
        demo_start = p
    if "Headings — use Heading 1" in text and bullet_numpr is None:
        numpr = p.find(qn("w:pPr") + "/" + qn("w:numPr"))
        if numpr is not None:
            bullet_numpr = copy.deepcopy(numpr)
    if "{First step or point.}" in text and ordered_numpr is None:
        numpr = p.find(qn("w:pPr") + "/" + qn("w:numPr"))
        if numpr is not None:
            ordered_numpr = copy.deepcopy(numpr)
    if text.startswith("Figure ") and "—" in text and caption_ppr is None:
        ppr = p.find(qn("w:pPr"))
        caption_ppr = copy.deepcopy(ppr) if ppr is not None else None
        first_r = p.find(qn("w:r"))
        if first_r is not None:
            rpr = first_r.find(qn("w:rPr"))
            caption_rpr = copy.deepcopy(rpr) if rpr is not None else None

for r in body.iter(qn("w:r")):
    rtext = "".join(t.text or "" for t in r.findall(qn("w:t")))
    rpr = r.find(qn("w:rPr"))
    if rpr is None:
        continue
    if "(MUST)" in rtext and must_rpr is None:
        must_rpr = copy.deepcopy(rpr)
    if "(SHOULD)" in rtext and should_rpr is None:
        should_rpr = copy.deepcopy(rpr)
    if "(MAY)" in rtext and may_rpr is None:
        may_rpr = copy.deepcopy(rpr)

for tbl in body.findall(qn("w:tbl")):
    if "{Column A}" in "".join(t.text or "" for t in tbl.iter(qn("w:t"))):
        exemplar_tbl = tbl
        break
if exemplar_tbl is None:
    # fall back: last table before demo removal that has a shaded header row
    tables = body.findall(qn("w:tbl"))
    exemplar_tbl = tables[-1] if tables else None

if demo_start is None:
    sys.exit("demo marker paragraph not found")
if exemplar_tbl is None:
    sys.exit("exemplar table not found")

ex_tblpr = copy.deepcopy(exemplar_tbl.find(qn("w:tblPr")))
ex_rows = exemplar_tbl.findall(qn("w:tr"))
ex_header_tc = ex_rows[0].findall(qn("w:tc"))[0]
ex_body_tc = ex_rows[1].findall(qn("w:tc"))[0]
ex_zebra_tc = ex_rows[2].findall(qn("w:tc"))[0] if len(ex_rows) > 2 else ex_body_tc
ex_header_trpr = ex_rows[0].find(qn("w:trPr"))


def harvested_cell_bits(tc):
    """Return (tcPr, pPr, rPr) deep copies from an exemplar cell."""
    tcpr = tc.find(qn("w:tcPr"))
    p = tc.find(qn("w:p"))
    ppr = p.find(qn("w:pPr")) if p is not None else None
    r = p.find(qn("w:r")) if p is not None else None
    rpr = r.find(qn("w:rPr")) if r is not None else None
    return (
        copy.deepcopy(tcpr) if tcpr is not None else None,
        copy.deepcopy(ppr) if ppr is not None else None,
        copy.deepcopy(rpr) if rpr is not None else None,
    )


HEADER_BITS = harvested_cell_bits(ex_header_tc)
BODY_BITS = harvested_cell_bits(ex_body_tc)
ZEBRA_BITS = harvested_cell_bits(ex_zebra_tc)

# --------------------------------------------------------------------------
# 2. Remove the demo pages (keep the final sectPr)
# --------------------------------------------------------------------------
children = list(body)
start_index = children.index(demo_start)
removed = 0
for el in children[start_index:]:
    if el.tag == qn("w:sectPr"):
        continue
    body.remove(el)
    removed += 1
print(f"removed {removed} demo elements")

sectpr = body.find(qn("w:sectPr"))


def append_element(el):
    """Insert an element at the end of the body but before the sectPr."""
    if sectpr is not None:
        sectpr.addprevious(el)
    else:
        body.append(el)


# --------------------------------------------------------------------------
# 3. Markdown rendering helpers
# --------------------------------------------------------------------------

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))")


def make_paragraph(style_id=None):
    p = OxmlElement("w:p")
    append_element(p)
    from docx.text.paragraph import Paragraph

    para = Paragraph(p, doc)
    if style_id:
        ppr = p.get_or_add_pPr()
        pstyle = OxmlElement("w:pStyle")
        pstyle.set(qn("w:val"), style_id)
        ppr.append(pstyle)
    return para


def add_hyperlink(paragraph, text, url):
    r_id = doc.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def add_inline(paragraph, text, base_rpr=None):
    """Render markdown inline formatting into runs."""
    for token in INLINE_RE.split(text):
        if not token:
            continue
        m_link = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
        if m_link:
            add_hyperlink(paragraph, m_link.group(1), m_link.group(2))
            continue
        bold = italic = code = False
        content = token
        if token.startswith("**") and token.endswith("**"):
            bold, content = True, token[2:-2]
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            italic, content = True, token[1:-1]
        elif token.startswith("`") and token.endswith("`"):
            code, content = True, token[1:-1]
        run = paragraph.add_run(content)
        if base_rpr is not None and run._r.find(qn("w:rPr")) is None:
            run._r.insert(0, copy.deepcopy(base_rpr))
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if code:
            run.font.name = "JetBrains Mono"
            run.font.size = None
        # requirement keywords take exemplar colouring
        if content == "(MUST)" and must_rpr is not None:
            _swap_rpr(run, must_rpr, keep_bold=True)
        elif content == "(SHOULD)" and should_rpr is not None:
            _swap_rpr(run, should_rpr, keep_bold=True)
        elif content == "(MAY)" and may_rpr is not None:
            _swap_rpr(run, may_rpr, keep_bold=True)
        elif content == "(SHOULD)":
            run.font.color.rgb = AMBER_TEXT
        elif content == "(MAY)":
            run.font.color.rgb = GREY


def _swap_rpr(run, rpr, keep_bold=False):
    old = run._r.find(qn("w:rPr"))
    if old is not None:
        run._r.remove(old)
    run._r.insert(0, copy.deepcopy(rpr))
    if keep_bold:
        run.bold = True


def add_heading(text, level):
    para = make_paragraph(f"Heading{level}")
    add_inline(para, text)


def add_body_paragraph(text):
    para = make_paragraph()
    add_inline(para, text)


def add_list_item(text, numpr, indent_fallback):
    para = make_paragraph("ListParagraph")
    ppr = para._p.get_or_add_pPr()
    if numpr is not None:
        ppr.append(copy.deepcopy(numpr))
    else:
        add_inline(para, ("• " if indent_fallback == "bullet" else "") + text)
        return
    add_inline(para, text)


def add_image_with_caption(path, caption):
    img_path = (BASE / path).resolve()
    with Image.open(img_path) as im:
        w_px, h_px = im.size
    scale = min(MAX_W_IN / (w_px / 220), MAX_H_IN / (h_px / 220), 1.0)
    width = Inches(w_px / 220 * scale)
    para = make_paragraph()
    para.alignment = 1  # centre
    run = para.add_run()
    run.add_picture(str(img_path), width=width)
    cap = make_paragraph()
    if caption_ppr is not None:
        old = cap._p.find(qn("w:pPr"))
        if old is not None:
            cap._p.remove(old)
        cap._p.insert(0, copy.deepcopy(caption_ppr))
    add_inline(cap, caption, base_rpr=caption_rpr)


def styled_cell(tc, bits, text, header=False):
    tcpr, ppr, rpr = bits
    for old in tc.findall(qn("w:tcPr")):
        tc.remove(old)
    if tcpr is not None:
        tc.insert(0, copy.deepcopy(tcpr))
    # first paragraph
    p = tc.find(qn("w:p"))
    for extra in tc.findall(qn("w:p"))[1:]:
        tc.remove(extra)
    for child in list(p):
        p.remove(child)
    if ppr is not None:
        p.append(copy.deepcopy(ppr))
    from docx.text.paragraph import Paragraph

    para = Paragraph(p, doc)
    add_inline(para, text, base_rpr=rpr)


def add_table(header_cells, rows):
    tbl = OxmlElement("w:tbl")
    tbl.append(copy.deepcopy(ex_tblpr))
    grid = OxmlElement("w:tblGrid")
    for _ in header_cells:
        grid.append(OxmlElement("w:gridCol"))
    tbl.append(grid)

    def make_row(cells, bits, header=False):
        tr = OxmlElement("w:tr")
        if header and ex_header_trpr is not None:
            tr.append(copy.deepcopy(ex_header_trpr))
        for cell_text in cells:
            tc = OxmlElement("w:tc")
            tcp = OxmlElement("w:tcPr")
            tc.append(tcp)
            p = OxmlElement("w:p")
            tc.append(p)
            tr.append(tc)
            styled_cell(tc, bits, cell_text, header=header)
        return tr

    tbl.append(make_row(header_cells, HEADER_BITS, header=True))
    for i, row in enumerate(rows):
        row = row + [""] * (len(header_cells) - len(row))
        bits = ZEBRA_BITS if i % 2 == 1 else BODY_BITS
        tbl.append(make_row(row[: len(header_cells)], bits))
    append_element(tbl)
    # spacer paragraph after table (Word requires a paragraph between tables)
    make_paragraph()


# --------------------------------------------------------------------------
# 4. Parse the markdown and emit
# --------------------------------------------------------------------------
lines = []
for raw_line in SRS_MD.read_text().splitlines():
    m_inc = re.match(r"\s*<!--\s*include:\s*(\S+)\s*-->\s*$", raw_line)
    if m_inc:
        lines.extend((BASE / m_inc.group(1)).read_text().splitlines())
    else:
        lines.append(raw_line)
i = 0
para_buf = []


def add_page_break():
    para = make_paragraph()
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    para._p.append(run)


def flush_paragraph():
    global para_buf
    if para_buf:
        add_body_paragraph(" ".join(para_buf))
        para_buf = []


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    if stripped.startswith("#"):
        flush_paragraph()
        level = len(stripped) - len(stripped.lstrip("#"))
        add_heading(stripped.lstrip("#").strip(), min(level, 3))
    elif stripped.startswith("!["):
        flush_paragraph()
        m = re.match(r"!\[(.*?)\]\((.*?)\)", stripped, re.S)
        # caption may wrap over lines until the closing )
        full = stripped
        while not re.match(r"!\[(.*)\]\((.*)\)\s*$", full, re.S) and i + 1 < len(lines):
            i += 1
            full += " " + lines[i].strip()
        m = re.match(r"!\[(.*)\]\((.*)\)\s*$", full, re.S)
        add_image_with_caption(m.group(2), m.group(1))
    elif stripped.startswith("|"):
        flush_paragraph()
        header = split_row(stripped)
        i += 1  # separator row
        rows = []
        while i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            i += 1
            rows.append(split_row(lines[i]))
        add_table(header, rows)
    elif stripped.startswith("- "):
        flush_paragraph()
        item = stripped[2:]
        while (
            i + 1 < len(lines)
            and lines[i + 1].startswith("  ")
            and lines[i + 1].strip()
            and not lines[i + 1].lstrip().startswith("- ")
        ):
            i += 1
            item += " " + lines[i].strip()
        add_list_item(item, bullet_numpr, "bullet")
    elif re.match(r"^\d+\.\s", stripped):
        flush_paragraph()
        item = re.sub(r"^\d+\.\s+", "", stripped)
        while i + 1 < len(lines) and lines[i + 1].startswith("   ") and lines[i + 1].strip():
            i += 1
            item += " " + lines[i].strip()
        add_list_item(item, ordered_numpr, "ordered")
    elif stripped == "<!-- pagebreak -->":
        flush_paragraph()
        add_page_break()
    elif stripped.startswith("<!--"):
        flush_paragraph()  # other comments are ignored
    elif stripped == "---":
        flush_paragraph()
    elif stripped == "":
        flush_paragraph()
    else:
        para_buf.append(stripped)
    i += 1
flush_paragraph()

doc.save(str(OUT))
print(f"document written: {OUT}")


# --------------------------------------------------------------------------
# 5. Zip-level post pass: cover artwork + placeholder text
# --------------------------------------------------------------------------
def png_size(data):
    if data[:8] != b"\x89PNG\r\n\x1a\n":
        return (0, 0)
    return struct.unpack(">II", data[16:24])


parts = {}
with zipfile.ZipFile(OUT) as z:
    for name in z.namelist():
        parts[name] = z.read(name)

cover_bytes = COVER.read_bytes()
target_size = png_size(cover_bytes)
swapped = False
for name, data in parts.items():
    if name.startswith("word/media/") and png_size(data) == target_size:
        parts[name] = cover_bytes
        swapped = True
        print(f"cover swapped into {name}")
if not swapped:
    sys.exit("cover artwork not found by dimensions")

for part in ("word/document.xml", "word/header1.xml", "word/footer1.xml"):
    if part in parts:
        text = parts[part].decode()
        for old, new in REPLACEMENTS.items():
            text = text.replace(old, new)
        parts[part] = text.encode()

with zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED) as z:
    for name, data in parts.items():
        z.writestr(name, data)
print(f"final document: {OUT} ({OUT.stat().st_size // 1024} KiB)")
